import pandas as pd
# from datetime import datetime
from docx import Document

df = pd.read_excel("pst_art.xlsx")
dct = df.to_dict(orient='index')

j_ru = 'Петербургская социология сегодня'
j_en = 'St. Petersburg Sociology Today'
year = '2024'
issue = 25


def date_frmt(ts):
    if type(ts) is pd.Timestamp:
        return ts.strftime("%d.%m.%Y")
    else:
        return ts


for i in dct:
    dct[i]['received'] = date_frmt(dct[i]['received'])
    dct[i]['revised'] = date_frmt(dct[i]['revised'])
    dct[i]['accepted'] = date_frmt(dct[i]['accepted'])

art_nums = [v['art_num'] for k, v in dct.items() if v['issue'] == issue]
art_nums.sort()

items_lst = [v for i in art_nums for k, v in dct.items() if i == v['art_num']]


def add_content(document, content):
    section = ''
    for c in content:
        if c[0] != section:
            section = c[0]
            document.add_paragraph(section)
        p = document.add_paragraph('')
        p.add_run(c[1]).italic = True
        p.add_run(f" {c[2]}")
    document.add_page_break()


def get_fio(rec, lang='ru'):
    if lang == 'ru':
        iof_l = rec['iof'].split(' ')
    elif lang == 'en':
        iof_l = rec['iof_en'].split(' ')
    f_name = iof_l[0]
    patr = iof_l[1]
    l_name = iof_l[2]
    fio = f"{l_name} {f_name[0]}. {patr[0]}."
    return fio


def get_cit(rec, fio, lang='ru'):
    cit = ''
    if lang == 'ru':
        cit = f"{fio} {rec['title']} // {j_ru}. {year}. № {issue}. С. {rec['pp']}. "
    elif lang == 'en':
        cit = f"{fio} {rec['title_en']}. {j_en}. {year}. No {issue}. P. {rec['pp']}. "
    cit += f"DOI: {rec['doi']}; EDN: {rec['edn']}"
    return cit


def get_receiving(rec, lang='ru'):
    receiving = ''
    if lang == 'ru':
        receiving = f"Статья поступила в редакцию: {rec['received']}; "
        if rec['revised']:
            receiving += f"поступила после рецензирования и доработки: {rec['revised']}; "
        receiving += f"принята к публикации: {rec['accepted']}."
    elif lang == 'en':
        receiving = f"Received: {rec['received']}; "
        if rec['revised']:
            receiving += f"revised after review: {rec['revised']}; "
        receiving += f"accepted for publication: {rec['accepted']}."
    return receiving


def author_info(rec, lang='ru'):
    if lang == 'ru':
        title = 'Сведения об авторе'
        iof_l = rec['iof'].split(' ')
        a_name = f"{iof_l[2]} {iof_l[0]} {iof_l[1]}"
        aff = rec['aff']
    else:
        title = 'Information about the author'
        iof_l = rec['iof_en'].split(' ')
        a_name = f"{iof_l[2]} {iof_l[0]} {iof_l[1]}"
        aff = rec['aff_en']
    eml = rec['email']
    return title, a_name, aff, eml


def add_art(document, rec, cit_ru, cit_en, receiving_ru, receiving_en):
    document.add_heading(rec['section'].upper(), level=2)
    document.add_paragraph(f"DOI: {rec['doi']}")
    document.add_paragraph(f"EDN: {rec['edn']}")
    document.add_paragraph(f"УДК {rec['udc']}")

    document.add_paragraph(f"{rec['iof']}1")
    document.add_paragraph(f"1 {rec['aff']}")
    document.add_heading(rec['title'], level=3)
    p = document.add_paragraph('')
    p.add_run('Аннотация.').italic = True
    p.add_run(f" {rec['abstr']}")

    p = document.add_paragraph('')
    p.add_run('Ключевые слова:').italic = True
    p.add_run(f" {rec['kw']}")

    p = document.add_paragraph('')
    p.add_run('Ссылка для цитирования:').italic = True
    p.add_run(f" {cit_ru}")

    title, a_name, aff, eml = author_info(rec, lang='ru')
    document.add_heading(title, level=4)
    p = document.add_paragraph('')
    p.add_run(a_name).bold = True
    p.add_run(f", <...> {aff}. {eml}")
    document.add_paragraph(receiving_ru)

    document.add_paragraph(f"{rec['iof_en']}1")
    document.add_paragraph(f"1 {rec['aff_en']}")
    document.add_heading(rec['title_en'], level=3)
    p = document.add_paragraph('')
    p.add_run('Abstract.').italic = True
    p.add_run(f" {rec['abstr_en']}")

    p = document.add_paragraph('')
    p.add_run('Keywords:').italic = True
    p.add_run(f" {rec['kw_en']}")

    p = document.add_paragraph('')
    p.add_run('For citation:').italic = True
    p.add_run(f" {cit_en}")

    title, a_name, aff, eml = author_info(rec, lang='en')
    document.add_heading(title, level=4)
    p = document.add_paragraph('')
    p.add_run(a_name).bold = True
    p.add_run(f", <...> {aff}. {eml}")
    document.add_paragraph(receiving_en)

    document.add_page_break()



doc = Document()

doc.add_heading('СОДЕРЖАНИЕ', level=2)
content_ru = []
for i in items_lst:
    fio = get_fio(i, lang='ru')
    content_ru.append((i['section'], fio, i['title']))
add_content(doc, content_ru)

doc.add_heading('CONTENTS', level=2)
content_en = []
for i in items_lst:
    fio = get_fio(i, lang='en')
    content_en.append((i['section_en'], fio, i['title_en']))
add_content(doc, content_en)

for i in items_lst:
    fio_ru = get_fio(i, lang='ru')
    fio_en = get_fio(i, lang='en')
    cit_ru = get_cit(i, fio_ru, lang='ru')
    cit_en = get_cit(i, fio_en, lang='en')
    receiving_ru = get_receiving(i, lang='ru')
    receiving_en = get_receiving(i, lang='en')
    add_art(doc, i, cit_ru, cit_en, receiving_ru, receiving_en)

doc.save(f'pst_{issue}.docx')
